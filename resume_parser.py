import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
import json

load_dotenv()

my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("Where is api key?")

client = Groq(api_key=my_api_key)

model= "openai/gpt-oss-20b"
# role="user"

from pydantic import BaseModel
class Description(BaseModel):
    job_title:str
    required_skils: list[str]
    preffered_skills: list[str]
    experience:int | None = None
    education:str | None = None

schema = Description.model_json_schema()
response_format={
    "type":"json_object"
}
system_prompt=f"""
return the output in json format strictly following this schema
{schema}
"""
message_system = {
    "role":"system",
    "content":system_prompt
}
def parse_job_description(job_description: str):

    schema = Description.model_json_schema()

    system_prompt = f"""
    Return the output in JSON format strictly following this schema:

    {schema}
    """

    message_system = {
        "role": "system",
        "content": system_prompt
    }

    prompt = f"""
    Extract the following information from the job description:

    - Job title
    - Required skills (return as a list of individual skills)
    - Preferred skills (return as a list)
    - Required experience (return only the number of years if mentioned)
    - Education (return a short phrase)

    Do not merge multiple skills into one sentence.

    JOB DESCRIPTION:
    {job_description}
    """

    message = {
        "role": "user",
        "content": prompt
    }

    messages = [message_system, message]

    response_format = {
        "type": "json_object"
    }

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format
    )

    answer = response.choices[0].message.content

    job_data = json.loads(answer)

    job = Description(**job_data)

    return job


# print(job.job_title)
# print(job.required_skils)
# print(job.preffered_skills)

# for skills in job.required_skils:
#     print(skills)

# for skills in job.preffered_skills:
#     print(skills)


class MatchDetails(BaseModel):
    matching_skills : list[str] = []
    missing_skills : list[str] = []
    experience_requirement_met : bool
    verdict : str

class MatchResult(BaseModel):
    score: float
    details: MatchDetails

class Experience(BaseModel):
    company:str | None = None
    role:str | None = None
    duration:str | None = None
    description:str | None = None
    skills_used : list[str] = []

class Resume(BaseModel):
    name:str | None = None
    email:str | None = None
    phone: str | None = None
    skills:list[str] = []
    education : list[str] = []
    projects : list[str] = []
    certifications : list[str] = []
    experiences:list[Experience] = []

resume_schema = Resume.model_json_schema()

def final_result(job, resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message = {
        "role":"user",
        "content":prompt
    }
    messages=[message]
    response_format={
        "type":"json_object"
    }
    response = client.chat.completions.create(model=model,messages=messages,response_format=response_format)
    raw_data = response.choices[0].message.content
    data = json.loads(raw_data)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    message_system = {
        "role" : "system",
        "content" : system_prompt
    }

    prompt = f"""
    Parse the following resume

    {resume_text}
    """
    message = {
        "role" : "user",
        "content" : prompt
    }

    messages = [message_system,message]
    response_format = {
        "type" : "json_object"
    }

    response = client.chat.completions.create(model=model,messages=messages,response_format=response_format)
    answer = response.choices[0].message.content
    raw_output = answer
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume



from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text();

        if page_text:
            text+=page_text + "\n"

    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    
    for paragraph in document.paragraphs:
        # paragraph_text = paragraphs.extract_text();   #! okay here i cant do this extract text because here the text is in continous format which can have a lot of spaces and all so we need to remove those spaces and make the output clear this is why we use strip.
        if paragraph.text.strip():
            text+=paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None
    

def analyze_resume(file_path, job_description):

    # Convert the raw job description into structured job data
    job = parse_job_description(job_description)

    # Read the resume
    resume_text = read_resume(file_path)

    # Parse the resume
    parsed_resume = parse_resume(resume_text)

    # Compare the resume with the job description
    result = final_result(job, parsed_resume)

    return {
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    }


if __name__ == "__main__":

    job_description = """
    We are looking for a Python Developer with experience in FastAPI,
    REST APIs, SQL, Git, and machine learning.
    """

    file_path = Path("resumes/Harshit_Khatri_resume.pdf")

    result = analyze_resume(
        file_path,
        job_description
    )

    print(result)




